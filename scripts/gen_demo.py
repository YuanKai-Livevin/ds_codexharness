#!/usr/bin/env python3
"""Generate a demo workspace with sample office files for trying the app."""
import os
import random

ROOT = os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "demo-workspace")
os.makedirs(ROOT, exist_ok=True)
random.seed(42)

# --- Excel: two quarters of sales data ---
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill

headers = ["日期", "区域", "产品", "销量", "单价", "销售额"]
regions = ["华东", "华北", "华南", "西南"]
products = ["笔记本", "显示器", "键盘", "鼠标"]

for year in (2025, 2026):
    wb = Workbook()
    ws = wb.active
    ws.title = f"销售数据{year}"
    ws.append(headers)
    for c in ws[1]:
        c.font = Font(bold=True)
        c.fill = PatternFill("solid", fgColor="D9E2F3")
    for month in range(1, 7):
        for _ in range(12):
            r = random.choice(regions)
            p = random.choice(products)
            qty = random.randint(20, 300)
            price = {"笔记本": 5999, "显示器": 1299, "键盘": 199, "鼠标": 89}[p]
            ws.append([f"{year}-{month:02d}-15", r, p, qty, price, qty * price])
    wb.save(os.path.join(ROOT, f"销售数据-{year}.xlsx"))

# --- Excel: employee roster ---
wb = Workbook()
ws = wb.active
ws.title = "员工名单"
ws.append(["工号", "姓名", "部门", "入职日期", "月薪"])
rows = [
    ["E001", "张伟", "销售部", "2019-03-12", 12000],
    ["E002", "李娜", "市场部", "2020-07-01", 11000],
    ["E003", "王强", "技术部", "2018-11-20", 15000],
    ["E004", "赵敏", "人事部", "2021-05-08", 9800],
    ["E005", "刘洋", "财务部", "2017-09-15", 13500],
    ["E006", "陈静", "销售部", "2022-01-10", 10500],
]
for r in rows:
    ws.append(r)
wb.save(os.path.join(ROOT, "员工名单.xlsx"))

# --- Word: meeting minutes ---
from docx import Document

doc = Document()
doc.add_heading("季度销售总结会议纪要", level=1)
doc.add_paragraph("会议时间：2026年3月30日 14:00")
doc.add_paragraph("参会人员：张伟、李娜、王强、赵敏")
doc.add_heading("一、会议议题", level=2)
for t in ["1. 上季度销售数据回顾", "2. 新产品推广计划", "3. 人员分工调整"]:
    doc.add_paragraph(t)
doc.add_heading("二、会议决议", level=2)
doc.add_paragraph("1. 下季度重点推广「笔记本」产品线；")
doc.add_paragraph("2. 华南区增加 2 名销售人员；")
doc.add_paragraph("3. 每月 5 日前提交销售汇总表。")
doc.save(os.path.join(ROOT, "季度会议纪要.docx"))

print("demo workspace ready:", ROOT)
for f in sorted(os.listdir(ROOT)):
    print(" ", f)
