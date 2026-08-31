# -*- coding: utf-8 -*-
"""Word / PDF / 图片 / 文件工具（T0-04：全部走安全写入事务）。"""

import os

from .common import (
    ToolError, check_conflicts, ensure_dir, list_paths, out, require, resolve, safe_write,
    validate_output,
)


# ---------- Word ----------

def word_fill_template(params, root, dry_run=False):
    p = require(["input", "output", "values"], params, ["overwrite"])
    src = resolve(root, p["input"])
    if not os.path.isfile(src):
        raise ToolError("FILE_NOT_FOUND", f"模板不存在: {p['input']}")
    out_path = resolve(root, p["output"])
    overwrite = bool(p.get("overwrite", False))
    values = p["values"]
    if not isinstance(values, dict):
        raise ToolError("SCHEMA_ERROR", "values 必须是对象 {占位符: 值}")
    if dry_run:
        return out(True, f"预演：填充 {len(values)} 个占位符 → {out_path}",
                   outputs=[out_path], extra={"placeholder_count": len(values)})
    import docx
    doc = docx.Document(src)
    missing = []
    replaced = 0

    def fill_text(t):
        nonlocal replaced, missing
        if not t:
            return t
        new = t
        for k, v in values.items():
            token = "{{" + k + "}}"
            if token in new:
                new = new.replace(token, str(v))
                replaced += 1
        # 统计未替换占位符
        import re
        for m in re.findall(r"\{\{\s*([^}]+?)\s*\}\}", new):
            if m not in missing:
                missing.append(m)
        return new

    for para in doc.paragraphs:
        for run in para.runs:
            run.text = fill_text(run.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.text = fill_text(run.text)

    def wf(tmp):
        doc.save(tmp)

    info = safe_write(root, out_path, "docx", overwrite, wf)
    warn = [f"未找到占位符值: {m}" for m in missing] if missing else None
    return out(True, f"模板填充完成：替换 {replaced} 处，未替换 {len(missing)} 个占位符 → {out_path}",
               outputs=[out_path], warnings=warn,
               extra={"replaced": replaced, "missing": missing,
                      "overwrite": overwrite, "old_hash": info["old_hash"],
                      "new_hash": info["new_hash"], "backup": info["backup"]})


def word_extract_text(params, root, dry_run=False):
    p = require(["input"], params, ["output_txt", "overwrite"])
    src = resolve(root, p["input"])
    if not os.path.isfile(src):
        raise ToolError("FILE_NOT_FOUND", f"文件不存在: {p['input']}")
    out_txt = resolve(root, p["output_txt"]) if p.get("output_txt") else None
    if dry_run:
        return out(True, f"预演：提取正文文本 {src}")
    import docx
    doc = docx.Document(src)
    paras = [para.text for para in doc.paragraphs if para.text.strip()]
    text = "\n".join(paras)
    outputs = []
    if out_txt:
        overwrite = bool(p.get("overwrite", False))

        def wf(tmp):
            with open(tmp, "w", encoding="utf-8") as fh:
                fh.write(text)

        info = safe_write(root, out_txt, "any", overwrite, wf)
        outputs.append(out_txt)
        return out(True, f"提取完成：{len(paras)} 个段落，{len(text)} 字符 → {out_txt}",
                   outputs=outputs,
                   extra={"paragraphs": len(paras), "chars": len(text),
                          "overwrite": overwrite, "old_hash": info["old_hash"],
                          "new_hash": info["new_hash"], "backup": info["backup"]})
    return out(True, f"提取完成：{len(paras)} 个段落，{len(text)} 字符",
               outputs=outputs, extra={"paragraphs": len(paras), "chars": len(text)})


# ---------- PDF ----------

def pdf_merge(params, root, dry_run=False):
    p = require(["inputs", "output"], params, ["overwrite"])
    inputs = list_paths(root, p["inputs"], exts={".pdf"})
    if not inputs:
        raise ToolError("EMPTY_INPUT", "没有可合并的 PDF")
    out_path = resolve(root, p["output"])
    if os.path.splitext(out_path)[1].lower() != ".pdf":
        raise ToolError("SCHEMA_ERROR", "输出必须是 .pdf")
    overwrite = bool(p.get("overwrite", False))
    if dry_run:
        return out(True, f"预演：合并 {len(inputs)} 个 PDF → {out_path}", outputs=[out_path])
    from pypdf import PdfReader, PdfWriter
    writer = PdfWriter()
    total = 0
    for f in inputs:
        r = PdfReader(f)
        for page in r.pages:
            writer.add_page(page)
        total += len(r.pages)

    def wf(tmp):
        with open(tmp, "wb") as fh:
            writer.write(fh)

    info = safe_write(root, out_path, "pdf", overwrite, wf)
    return out(True, f"PDF 合并完成：{len(inputs)} 个文件共 {total} 页 → {out_path}",
               outputs=[out_path],
               extra={"files": len(inputs), "pages": total,
                      "overwrite": overwrite, "old_hash": info["old_hash"],
                      "new_hash": info["new_hash"], "backup": info["backup"]})


def pdf_split(params, root, dry_run=False):
    p = require(["input", "output_dir"], params, ["ranges", "pages", "overwrite"])
    src = resolve(root, p["input"])
    if not os.path.isfile(src):
        raise ToolError("FILE_NOT_FOUND", f"文件不存在: {p['input']}")
    out_dir = resolve(root, p["output_dir"])
    overwrite = bool(p.get("overwrite", False))
    ranges = p.get("ranges") or []
    pages = p.get("pages") or []
    if not ranges and not pages:
        ranges = [[1, 1 << 30]]  # 默认每页一个文件
    if dry_run:
        return out(True, f"预演：拆分 {src} → {out_dir}", outputs=[out_dir])
    from pypdf import PdfReader, PdfWriter
    reader = PdfReader(src)
    total = len(reader.pages)
    ensure_dir(out_dir)
    base = os.path.splitext(os.path.basename(src))[0]
    # 预计算输出计划
    plan = []  # (dst, writer_build)
    idx = 0
    if pages:
        for pg in pages:
            if not (1 <= pg <= total):
                raise ToolError("SCHEMA_ERROR", f"页码越界: {pg}（共 {total} 页）")
            idx += 1
            dst = os.path.join(out_dir, f"{base}-p{pg:03d}.pdf")
            plan.append((dst, [pg - 1]))
    else:
        for lo, hi in ranges:
            lo = max(1, int(lo))
            hi = min(total, int(hi))
            if lo > total:
                raise ToolError("SCHEMA_ERROR", f"起始页越界: {lo}（共 {total} 页）")
            idx += 1
            dst = os.path.join(out_dir, f"{base}-part{idx:03d}.pdf")
            plan.append((dst, list(range(lo - 1, hi))))
    # 批量预检：任一输出已存在且未批准 → 整体不写
    check_conflicts([d for d, _ in plan], overwrite)
    outputs = []
    for dst, pg_list in plan:
        w = PdfWriter()
        for pgi in pg_list:
            w.add_page(reader.pages[pgi])

        def wf(tmp, w=w):
            with open(tmp, "wb") as fh:
                w.write(fh)

        safe_write(root, dst, "pdf", overwrite, wf)
        outputs.append(dst)
    return out(True, f"PDF 拆分完成：共 {len(outputs)} 个文件 → {out_dir}",
               outputs=outputs,
               extra={"parts": len(outputs), "overwrite": overwrite})


def pdf_text(params, root, dry_run=False):
    p = require(["input"], params, ["output_txt", "first_page", "last_page", "overwrite"])
    src = resolve(root, p["input"])
    if not os.path.isfile(src):
        raise ToolError("FILE_NOT_FOUND", f"文件不存在: {p['input']}")
    out_txt = resolve(root, p["output_txt"]) if p.get("output_txt") else None
    if dry_run:
        return out(True, f"预演：提取文本 {src}")
    from pypdf import PdfReader
    reader = PdfReader(src)
    total = len(reader.pages)
    lo = int(p.get("first_page") or 1)
    hi = int(p.get("last_page") or total)
    lo, hi = max(1, lo), min(total, hi)
    parts = []
    for i in range(lo - 1, hi):
        parts.append(reader.pages[i].extract_text() or "")
    text = "\n\n".join(parts)
    outputs = []
    if out_txt:
        overwrite = bool(p.get("overwrite", False))

        def wf(tmp):
            with open(tmp, "w", encoding="utf-8") as fh:
                fh.write(text)

        info = safe_write(root, out_txt, "any", overwrite, wf)
        outputs.append(out_txt)
        return out(True, f"PDF 文本提取完成：{hi - lo + 1} 页，{len(text)} 字符 → {out_txt}",
                   outputs=outputs,
                   extra={"pages": hi - lo + 1, "chars": len(text),
                          "overwrite": overwrite, "old_hash": info["old_hash"],
                          "new_hash": info["new_hash"], "backup": info["backup"]})
    return out(True, f"PDF 文本提取完成：{hi - lo + 1} 页，{len(text)} 字符",
               outputs=outputs, extra={"pages": hi - lo + 1, "chars": len(text)})


# ---------- 图片 ----------

def _collect_images(root, inputs):
    exts = {".png", ".jpg", ".jpeg", ".bmp", ".gif", ".webp", ".tiff"}
    if inputs:
        return list_paths(root, inputs, exts=exts)
    # 默认：工作区根目录下的图片
    return list_paths(root, ["."], exts=exts)


def _img_save(tmp, img, fmt=None):
    if fmt:
        img.save(tmp, "JPEG" if fmt == "jpg" else fmt)
    else:
        img.save(tmp)
    img.close()


def img_resize(params, root, dry_run=False):
    p = require(["output_dir"], params, ["inputs", "width", "height", "percent", "overwrite"])
    files = _collect_images(root, p.get("inputs"))
    if not files:
        raise ToolError("EMPTY_INPUT", "没有找到图片")
    out_dir = resolve(root, p["output_dir"])
    overwrite = bool(p.get("overwrite", False))
    width, height, percent = p.get("width"), p.get("height"), p.get("percent")
    if width is None and height is None and percent is None:
        raise ToolError("SCHEMA_ERROR", "至少提供 width / height / percent 之一")
    if width is not None:
        width = int(width)
    if height is not None:
        height = int(height)
    if percent is not None:
        percent = float(percent)
    if dry_run:
        return out(True, f"预演：缩放 {len(files)} 张图片 → {out_dir}", outputs=[out_dir])
    from PIL import Image
    ensure_dir(out_dir)
    # 预计算输出计划 + 批量预检（多目录同名场景：任一冲突整体不写）
    plan = []
    for f in files:
        img = Image.open(f)
        ow, oh = img.size
        if percent:
            nw, nh = max(1, int(ow * percent)), max(1, int(oh * percent))
        else:
            nw = width or int(ow * (height / oh))
            nh = height or int(oh * (width / ow))
        nw, nh = max(1, nw), max(1, nh)
        if img.mode in ("P", "RGBA") and os.path.splitext(f)[1].lower() in (".jpg", ".jpeg"):
            img = img.convert("RGB")
        dst = os.path.join(out_dir, os.path.basename(f))
        plan.append((dst, img.resize((nw, nh), Image.LANCZOS)))
    check_conflicts([d for d, _ in plan], overwrite)
    outputs = []
    for dst, resized in plan:
        safe_write(root, dst, "image", overwrite, lambda tmp, im=resized: _img_save(tmp, im))
        outputs.append(dst)
    return out(True, f"图片缩放完成：{len(outputs)} 张 → {out_dir}",
               outputs=outputs,
               extra={"count": len(outputs), "overwrite": overwrite})


def img_convert(params, root, dry_run=False):
    p = require(["output_dir", "format"], params, ["inputs", "overwrite"])
    fmt = (p["format"] or "").lower().lstrip(".")
    if fmt not in ("png", "jpg", "jpeg", "webp"):
        raise ToolError("SCHEMA_ERROR", "format 必须是 png|jpg|webp")
    if fmt == "jpeg":
        fmt = "jpg"
    files = _collect_images(root, p.get("inputs"))
    if not files:
        raise ToolError("EMPTY_INPUT", "没有找到图片")
    out_dir = resolve(root, p["output_dir"])
    overwrite = bool(p.get("overwrite", False))
    if dry_run:
        return out(True, f"预演：转换 {len(files)} 张图片为 {fmt} → {out_dir}", outputs=[out_dir])
    from PIL import Image
    ensure_dir(out_dir)
    plan = []
    for f in files:
        img = Image.open(f)
        if fmt in ("jpg",) and img.mode in ("P", "RGBA"):
            img = img.convert("RGB")
        base = os.path.splitext(os.path.basename(f))[0]
        dst = os.path.join(out_dir, base + "." + fmt)
        plan.append((dst, img))
    check_conflicts([d for d, _ in plan], overwrite)
    outputs = []
    for dst, img in plan:
        safe_write(root, dst, "image", overwrite, lambda tmp, im=img, f=fmt: _img_save(tmp, im, f))
        outputs.append(dst)
    return out(True, f"图片转换完成：{len(outputs)} 张 → {out_dir}",
               outputs=outputs,
               extra={"count": len(outputs), "format": fmt, "overwrite": overwrite})


# ---------- 文件 ----------

def _sha256(path):
    import hashlib
    h = hashlib.sha256()
    with open(path, "rb") as fh:
        for chunk in iter(lambda: fh.read(65536), b""):
            h.update(chunk)
    return h.hexdigest()


def file_manifest(params, root, dry_run=False):
    p = require([], params, ["input_dir", "output_csv", "include_hash", "pattern", "overwrite"])
    base = resolve(root, p.get("input_dir") or ".")
    out_csv = resolve(root, p["output_csv"]) if p.get("output_csv") else None
    include_hash = bool(p.get("include_hash", True))
    pattern = p.get("pattern")
    if dry_run:
        return out(True, f"预演：生成文件清单 {base}" + (f" → {out_csv}" if out_csv else "（输出到标准输出）"))
    entries = []
    for cur, _dirs, files in os.walk(base):
        for f in sorted(files):
            fp = os.path.join(cur, f)
            rel = os.path.relpath(fp, base)
            if pattern and pattern.lower() not in f.lower():
                continue
            st = os.stat(fp)
            entries.append([rel, st.st_size, int(st.st_mtime), _sha256(fp) if include_hash else ""])
    if not entries:
        raise ToolError("EMPTY_INPUT", "目录中没有文件")
    header = ["path", "size", "mtime", "sha256"] if include_hash else ["path", "size", "mtime"]
    lines = [",".join(header)]
    for e in entries:
        lines.append(",".join([e[0], str(e[1]), str(e[2]), e[3] if include_hash else ""]))
    text = "\n".join(lines)
    outputs = []
    if out_csv:
        overwrite = bool(p.get("overwrite", False))

        def wf(tmp):
            with open(tmp, "w", encoding="utf-8-sig", newline="") as fh:
                fh.write(text)

        info = safe_write(root, out_csv, "csv", overwrite, wf)
        outputs.append(out_csv)
        return out(True, f"文件清单生成：{len(entries)} 个文件 → {out_csv}",
                   outputs=outputs,
                   extra={"count": len(entries), "total_bytes": sum(e[1] for e in entries),
                          "overwrite": overwrite, "old_hash": info["old_hash"],
                          "new_hash": info["new_hash"], "backup": info["backup"]})
    return out(True, f"文件清单生成：{len(entries)} 个文件",
               outputs=outputs,
               extra={"count": len(entries), "total_bytes": sum(e[1] for e in entries), "manifest": text[:4000]})


def file_rename(params, root, dry_run=False):
    """事务化批量重命名：预检全部 → 执行 → 中途失败反向恢复（回滚表）。"""
    p = require(["input_dir", "mapping"], params)
    base = resolve(root, p["input_dir"])
    mapping = p["mapping"]
    if not isinstance(mapping, list) or not mapping:
        raise ToolError("SCHEMA_ERROR", "mapping 必须是非空数组 [{from,to},...]")
    pairs = []
    for m in mapping:
        if not isinstance(m, dict) or "from" not in m or "to" not in m:
            raise ToolError("SCHEMA_ERROR", "mapping 元素必须是 {from,to}")
        src = resolve(base, m["from"])
        dst = resolve(base, m["to"])
        if not os.path.isfile(src):
            raise ToolError("FILE_NOT_FOUND", f"待重命名文件不存在: {m['from']}")
        if src == dst:
            raise ToolError("SCHEMA_ERROR", f"源与目标相同: {m['from']}")
        if os.path.exists(dst):
            raise ToolError("CONFLICT", f"目标已存在，拒绝覆盖: {m['to']}")
        pairs.append((src, dst))
    if dry_run:
        return out(True, f"预演：将重命名 {len(pairs)} 个文件", extra={"pairs": [[a, b] for a, b in pairs]})
    done = []
    for i, (src, dst) in enumerate(pairs):
        try:
            ensure_dir(os.path.dirname(dst) or base)
            os.rename(src, dst)
            done.append((src, dst))
        except OSError as e:
            # 中途失败：反向恢复已完成的改名
            rolled = []
            for rsrc, rdst in reversed(done):
                try:
                    os.rename(rdst, rsrc)
                    rolled.append([os.path.relpath(rdst, base), os.path.relpath(rsrc, base)])
                except OSError:
                    rolled.append(["恢复失败", os.path.relpath(rdst, base)])
            raise ToolError(
                "TOOL_FAILED",
                f"重命名第 {i + 1}/{len(pairs)} 项失败（{e}），已回滚 {len(rolled)} 项",
            )
    renamed = [[os.path.relpath(a, base), os.path.relpath(b, base)] for a, b in done]
    return out(True, f"重命名完成：{len(done)} 个文件", extra={"renamed": renamed})
