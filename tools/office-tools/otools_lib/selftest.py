# -*- coding: utf-8 -*-
"""内置单元测试：每个工具用临时夹具跑一遍，断言输出有效。"""

import csv
import os
import shutil
import sys
import tempfile

from . import excel_tools
from .office_tools import (
    file_manifest, file_rename, img_convert, img_resize, pdf_merge, pdf_split,
    pdf_text, word_extract_text, word_fill_template,
)


def _make_xlsx(path, rows):
    import openpyxl
    wb = openpyxl.Workbook()
    ws = wb.active
    for r in rows:
        ws.append(r)
    wb.save(path)
    wb.close()


def _make_docx(path, text):
    import docx
    doc = docx.Document()
    doc.add_paragraph(text)
    doc.save(path)


def _make_pdf(path, pages=2):
    from pypdf import PdfWriter
    w = PdfWriter()
    for _ in range(pages):
        w.add_blank_page(width=200, height=200)
    with open(path, "wb") as fh:
        w.write(fh)


def _make_img(path, mode="RGB", size=(40, 30)):
    from PIL import Image
    img = Image.new(mode, size, (200, 60, 60) if mode == "RGB" else (200, 60, 60, 255))
    img.save(path)


def run_selftest(root):
    """在临时目录跑全部工具测试；返回退出码（失败数）。"""
    work = tempfile.mkdtemp(prefix="otools-selftest-", dir=root)
    fails = []
    passed = []

    def check(name, result, extra_ok=None):
        ok = extra_ok(result) if extra_ok is not None else bool(result.get("ok"))
        (passed if ok else fails).append(name)
        print(("PASS  " if ok else "FAIL  ") + name + ("  | " + str(result.get("message", ""))[:90] if not ok else ""))

    try:
        # ---- Excel ----
        a = os.path.join(work, "a.xlsx")
        b = os.path.join(work, "b.xlsx")
        _make_xlsx(a, [["name", "qty"], ["苹果", 3], ["香蕉", 5]])
        _make_xlsx(b, [["name", "qty"], ["橙子", 2]])
        merged = os.path.join(work, "merged.xlsx")
        check("excel_merge", excel_tools.excel_merge(
            {"inputs": [a, b], "output": merged}, work))
        merged_dry = os.path.join(work, "merged-dry.xlsx")
        check("excel_merge dry-run 不写文件", excel_tools.excel_merge(
            {"inputs": [a, b], "output": merged_dry}, work, dry_run=True),
            extra_ok=lambda r: not os.path.exists(merged_dry))
        deduped = os.path.join(work, "deduped.xlsx")
        _make_xlsx(deduped, [["name", "qty"], ["苹果", 3], ["苹果", 3], ["香蕉", 5]])
        check("excel_dedupe", excel_tools.excel_dedupe(
            {"input": deduped, "output": os.path.join(work, "dd.xlsx"), "key_columns": ["name"]}, work),
            extra_ok=lambda r: r.get("rows_after") == 2)
        filtered = os.path.join(work, "filtered.xlsx")
        check("excel_filter", excel_tools.excel_filter(
            {"input": a, "output": filtered, "column": "qty", "op": "ge", "value": 4}, work),
            extra_ok=lambda r: r.get("rows_after") == 1)
        pivot = os.path.join(work, "pivot.xlsx")
        check("excel_pivot", excel_tools.excel_pivot(
            {"input": a, "output": pivot, "rows": ["name"], "values": "qty", "agg": "sum"}, work),
            extra_ok=lambda r: r.get("groups") == 2)
        check("excel_formula_check 正常", excel_tools.excel_formula_check({"input": a}, work))
        badf = os.path.join(work, "bad.xlsx")
        _make_xlsx(badf, [["a", "b"], ["=1/0", 2]])
        import openpyxl
        wb = openpyxl.load_workbook(badf)
        wb.active["A2"] = "#DIV/0!"
        wb.save(badf)
        wb.close()
        check("excel_formula_check 发现错误", excel_tools.excel_formula_check({"input": badf}, work),
              extra_ok=lambda r: r.get("ok") is False and r.get("error_count", 0) >= 1)

        # ---- Word ----
        tmpl = os.path.join(work, "tmpl.docx")
        _make_docx(tmpl, "尊敬的 {{name}}，您的工号是 {{id}}。")
        filled = os.path.join(work, "filled.docx")
        check("word_fill_template", word_fill_template(
            {"input": tmpl, "output": filled, "values": {"name": "张三", "id": "A001"}}, work))
        import docx as _d
        d = _d.Document(filled)
        check("word_fill_template 内容正确",
              {"ok": "张三" in "".join(p.text for p in d.paragraphs)})
        check("word_extract_text", word_extract_text(
            {"input": tmpl, "output_txt": os.path.join(work, "t.txt")}, work))

        # ---- PDF ----
        p1 = os.path.join(work, "p1.pdf")
        p2 = os.path.join(work, "p2.pdf")
        _make_pdf(p1, 2)
        _make_pdf(p2, 1)
        pm = os.path.join(work, "pm.pdf")
        check("pdf_merge", pdf_merge({"inputs": [p1, p2], "output": pm}, work),
              extra_ok=lambda r: r.get("pages") == 3)
        split_dir = os.path.join(work, "split")
        check("pdf_split", pdf_split(
            {"input": pm, "output_dir": split_dir, "ranges": [[1, 2], [3, 3]]}, work),
            extra_ok=lambda r: len(r.get("outputs", [])) == 2)
        check("pdf_text", pdf_text({"input": p1}, work))

        # ---- 图片 ----
        i1 = os.path.join(work, "i1.png")
        _make_img(i1, "RGBA")
        rs_dir = os.path.join(work, "rs")
        check("img_resize", img_resize(
            {"inputs": [i1], "output_dir": rs_dir, "width": 20}, work),
            extra_ok=lambda r: os.path.exists(os.path.join(rs_dir, "i1.png")))
        cv_dir = os.path.join(work, "cv")
        check("img_convert", img_convert(
            {"inputs": [i1], "output_dir": cv_dir, "format": "jpg"}, work),
            extra_ok=lambda r: os.path.exists(os.path.join(cv_dir, "i1.jpg")))

        # ---- 文件 ----
        fdir = os.path.join(work, "fdir")
        os.makedirs(fdir)
        with open(os.path.join(fdir, "报告.txt"), "w", encoding="utf-8") as fh:
            fh.write("hello")
        with open(os.path.join(fdir, "数据.csv"), "w", encoding="utf-8") as fh:
            fh.write("a,b\n1,2\n")
        mf = os.path.join(work, "manifest.csv")
        check("file_manifest", file_manifest(
            {"input_dir": fdir, "output_csv": mf}, work),
            extra_ok=lambda r: os.path.exists(mf))
        with open(mf, "r", encoding="utf-8-sig") as fh:
            rows = list(csv.reader(fh))
        check("file_manifest 内容", {"ok": len(rows) == 3})
        rn = os.path.join(fdir, "报告.txt")
        rn2 = os.path.join(fdir, "终稿.txt")
        check("file_rename", file_rename(
            {"input_dir": fdir, "mapping": [{"from": "报告.txt", "to": "终稿.txt"}]}, work),
            extra_ok=lambda r: os.path.exists(rn2) and not os.path.exists(rn))

        # ---- 权限边界 ----
        import json
        import subprocess
        script = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "otools.py")
        r = subprocess.run([sys.executable, script, "file_manifest", "--root", work,
                            "--params", json.dumps({"input_dir": os.path.dirname(work)})],
                           capture_output=True, text=True, encoding="utf-8")
        try:
            j = json.loads(r.stdout)
            check("路径越界被拒绝", {"ok": j.get("error_code") == "PERM_DENIED"})
        except Exception:  # noqa: BLE001
            check("路径越界被拒绝", {"ok": False})
        r2 = subprocess.run([sys.executable, script, "excel_merge", "--root", work,
                             "--params", json.dumps({"output": os.path.join(work, "x.xlsx")})],
                            capture_output=True, text=True, encoding="utf-8")
        try:
            j2 = json.loads(r2.stdout)
            check("缺参数被拒绝", {"ok": j2.get("error_code") == "SCHEMA_ERROR"})
        except Exception:  # noqa: BLE001
            check("缺参数被拒绝", {"ok": False})

        # ---- T0-04 安全写入事务 ----
        from .common import ToolError as _TE, safe_write, sha256_file, backup_dir_for
        # 1) 同名输出默认拒绝（CONFLICT）且原文件不变
        exist = os.path.join(work, "existing.xlsx")
        _make_xlsx(exist, [["name", "qty"], ["旧数据", 1]])
        orig_hash = sha256_file(exist)
        try:
            excel_tools.excel_filter(
                {"input": a, "output": exist, "column": "qty", "op": "ge", "value": 1}, work)
            check("同名输出默认拒绝", {"ok": False})
        except _TE as e:
            check("同名输出默认拒绝", {"ok": e.code == "CONFLICT" and sha256_file(exist) == orig_hash})
        # 2) 覆盖需显式批准：overwrite=true 成功、旧文件有备份、哈希不同
        over = os.path.join(work, "over.xlsx")
        _make_xlsx(over, [["name", "qty"], ["旧", 1]])
        old_hash = sha256_file(over)
        res = excel_tools.excel_filter(
            {"input": a, "output": over, "column": "qty", "op": "ge", "value": 1, "overwrite": True}, work)
        new_hash = sha256_file(over)
        check("覆盖需批准且备份", {"ok": res.get("backup") and old_hash != new_hash
                                      and res.get("old_hash") == old_hash and res.get("new_hash") == new_hash})
        check("备份文件存在", {"ok": res.get("backup") and os.path.exists(res["backup"])})
        # 3) 输出校验失败不影响旧文件（write_fn 写出无效文件 → 旧文件保留）
        keep = os.path.join(work, "keep.csv")
        with open(keep, "w", encoding="utf-8") as fh:
            fh.write("旧内容")
        try:
            safe_write(work, keep, "csv", True, lambda tmp: open(tmp, "w", encoding="utf-8").close())
            check("校验失败不影响旧文件", {"ok": False})
        except _TE as e:
            with open(keep, "r", encoding="utf-8") as fh:
                check("校验失败不影响旧文件", {"ok": e.code == "OUTPUT_INVALID" and fh.read() == "旧内容"})
        # 4) 批量冲突整体拒绝：一张图目标已存在 → 全部不写
        i2 = os.path.join(work, "i2.png")
        _make_img(i2, "RGB", (30, 30))
        bd = os.path.join(work, "batch")
        os.makedirs(bd)
        clash = os.path.join(bd, "i1.png")
        with open(clash, "wb") as fh:
            fh.write(b"occupied")
        try:
            img_resize({"inputs": [i1, i2], "output_dir": bd, "width": 10}, work)
            check("批量冲突整体拒绝", {"ok": False})
        except _TE as e:
            check("批量冲突整体拒绝",
                  {"ok": e.code == "CONFLICT" and not os.path.exists(os.path.join(bd, "i2.png"))})
        # 5) 重命名事务：正常执行 + 冲突预检拒绝
        check("rename 事务正常", file_rename(
            {"input_dir": fdir, "mapping": [{"from": "数据.csv", "to": "数据2.csv"}]}, work),
            extra_ok=lambda r: os.path.exists(os.path.join(fdir, "数据2.csv")))
        try:
            file_rename({"input_dir": fdir, "mapping": [{"from": "终稿.txt", "to": "数据2.csv"}]}, work)
            check("rename 冲突拒绝", {"ok": False})
        except _TE as e:
            check("rename 冲突拒绝", {"ok": e.code == "CONFLICT"})
    finally:
        shutil.rmtree(work, ignore_errors=True)

    print()
    print(f"== selftest: {len(passed)} 通过, {len(fails)} 失败 ==")
    if fails:
        print("失败项:", ", ".join(fails))
    return 1 if fails else 0
